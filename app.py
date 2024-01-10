from flask import Flask, render_template, request, session
import json
import requests
import config
import openai
from youtube_transcript_api import YouTubeTranscriptApi

from gtts import gTTS
from playsound import playsound

from googletrans import Translator
from fpdf import FPDF

import docx
from docx.shared import Pt


import sqlite3

translator = Translator()

app = Flask(__name__)
app.secret_key = config.secret_key

@app.route('/')
def home():
	return render_template('home.html')

@app.route("/signup")
def signup():
    
    
    name = request.args.get('username','')
    number = request.args.get('number','')
    email = request.args.get('email','')
    password = request.args.get('psw','')

    con = sqlite3.connect('signup.db')
    cur = con.cursor()
    cur.execute("insert into `detail` (`name`,`number`,`email`, `password`) VALUES (?, ?, ?, ?)",(name,number,email,password))
    con.commit()
    con.close()

    return render_template("signin.html")

@app.route("/signin")
def signin():

    mail1 = request.args.get('name','')
    password1 = request.args.get('psw','')
    con = sqlite3.connect('signup.db')
    cur = con.cursor()
    cur.execute("select `name`, `password` from detail where `name` = ? AND `password` = ?",(mail1,password1,))
    data = cur.fetchone()
    print(data)

    if data == None:
        return render_template("signup.html")    

    elif mail1 == 'admin' and password1 == 'admin':
        return render_template("index.html")

    elif mail1 == str(data[0]) and password1 == str(data[1]):
        return render_template("index.html")
    else:
        return render_template("signup.html")

@app.route('/index', methods=['GET', 'POST'])
def index():
    return render_template('index.html')


@app.route("/summarize", methods=["POST"])
def summarize():
    # Get YouTube API key
    youtube_api_key = config.youtube_API_key

    # Get video URL from request
    url = request.form["url"]

    # Extract video ID from URL
    video_id = url.split("v=")[1]

    # Make request to YouTube API to get video details
    api_url = f"https://www.googleapis.com/youtube/v3/videos?part=snippet&id={video_id}&key={youtube_api_key}"
    r = requests.get(api_url)
    video_data = r.json()

    # Extract video title and description
    title = video_data["items"][0]["snippet"]["title"]
    # description = video_data["items"][0]["snippet"]["description"]
    
    # Extract video transcript
    transcript = YouTubeTranscriptApi.get_transcript(video_id) 

    # Full transcript string
    transcript_str = ""
    for line in transcript:
        # Break if transcript exceeds 2048 tokens (max context length for text-davinci-002)
        if (len(transcript_str.strip()) + len(line["text"] + " ") + config.max_tokens) <= config.model_max_tokens:
            transcript_str += line["text"] + " "
        else:
            break

    # Initialize OpenAI API key
    openai.api_key = config.openai_API_key

    # Create prompt for GPT-3
    gpt3_prompt = (
        f"Summarize the video {title} with the following transcript: {transcript_str}"
    )

    completions = openai.Completion.create(
            engine=config.model,
            prompt=gpt3_prompt,
            max_tokens=config.max_tokens,
            n=1,
            stop=None,
            temperature=config.temperature,
        )
    
    '''try:
        # Generate summary with GPT-3
        completions = openai.Completion.create(
            engine=config.model,
            prompt=gpt3_prompt,
            max_tokens=config.max_tokens,
            n=1,
            stop=None,
            temperature=config.temperature,
        )

    except Exception as e:
        # Return error message as JSON response
        print("Exception raised when generating summary: ", e)
        return {"error": str(e)}'''

    # Extract summary from completions
    summary = completions.choices[0].text

    # Remove any extra newline '\n' characters
    summary = summary.replace("\n", "")
    session["a"] = summary

    return render_template('result.html',prediction = summary)

@app.route("/speech")
def speech():
    d=session.get("a",None) 
    myobj = gTTS(text=d, slow=False)
    myobj.save("text.mp3")
    playsound('text.mp3')
    return render_template("result.html",prediction = d)

@app.route("/translate",methods=['POST'])
def translate():
    lang = request.form['1']
    d=session.get("a",None) 
    translations = translator.translate(d, dest=lang)
    input_text = translations.text
    print(input_text)
    return render_template("result.html",prediction = input_text)

@app.route("/convert",methods=['POST'])
def convert():
    d=session.get("a",None)
    lang = request.form['1']
    print(lang)
     
    if str(lang) == str(0):

        file1 = open("summary.txt", "w")  # write mode
        file1.write(d)
        file1.close()
        
    
    elif str(lang) == str(1):
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size = 15)
        pdf.cell(200, 10, txt = "Summary:",ln = 1, align = 'C')
        pdf.cell(200, 10, txt = d,ln = 2, align = 'C')
        pdf.output("summary.pdf")
        

    elif str(lang) == str(2):
        doc = docx.Document()
        doc.add_heading('Summary', 0)
        para = doc.add_paragraph().add_run(d)
        para.font.size = Pt(12)
        doc.save('summary.docx')
        

    elif str(lang) == str(3):
        myobj = gTTS(text=d, slow=False)
        myobj.save("summary.mp3")
        
    return render_template("result.html",prediction = 'Write The Summary Successfull')

@app.route("/send",methods=['POST'])
def send():
    import pywhatkit
    d=session.get("a",None)
    pywhatkit.sendwhatmsg("+919133301034", d, 13, 30, 15, True, 2)
    return render_template("result.html",prediction = 'Message Sent')

@app.route('/reg')
def reg():
	return render_template('signup.html')

@app.route('/login')
def login():
	return render_template('signin.html')

if __name__ == '__main__':
    app.run()
